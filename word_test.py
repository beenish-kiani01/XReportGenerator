from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

document = Document()

document.add_heading("X/Twitter Report", level=1)

table = document.add_table(rows=1, cols=4)
table.style = "Table Grid"

headers = [
    "Serial No.",
    "Screenshot",
    "Post URL",
    "Handler ID"
]

for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

row = table.add_row().cells

row[0].text = "1"

paragraph = row[1].paragraphs[0]
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = paragraph.add_run()
run.add_picture("screenshots/test.png", width=Inches(2.5))

row[2].text = "https://x.com/example"
row[3].text = "@example"

for cell in row:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

document.save("X_Report_Test.docx")

print("Table report created successfully!")