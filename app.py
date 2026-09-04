import os
import re
import traceback
import requests
from html import escape

from flask import Flask, render_template, request, jsonify, send_file, session, redirect, url_for
from dotenv import load_dotenv

from playwright.sync_api import sync_playwright

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


# ============================================================
# CONFIGURATION
# ============================================================

load_dotenv()

TOKEN = os.getenv("X_BEARER_TOKEN")

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "change-this-secret-key")
LOGIN_PASSWORD = os.getenv("LOGIN_PASSWORD", "admin987")

SCREENSHOT_DIR = "screenshots"
REPORT_DIR = "reports"

os.makedirs(SCREENSHOT_DIR, exist_ok=True)
os.makedirs(REPORT_DIR, exist_ok=True)


# ============================================================
# SEARCH X POSTS
# ============================================================

def search_posts(query, number_of_posts):

    if not TOKEN:
        raise Exception("X_BEARER_TOKEN is missing.")

    url = "https://api.x.com/2/tweets/search/recent"

    headers = {
        "Authorization": f"Bearer {TOKEN}"
    }

    params = {
        "query": f"{query} -is:retweet",
        "max_results": max(
            10,
            min(int(number_of_posts), 100)
        ),
        "tweet.fields": "created_at,author_id",
        "expansions": "author_id",
        "user.fields": "username,profile_image_url"
    }

    response = requests.get(
        url,
        headers=headers,
        params=params,
        timeout=30
    )

    print("X API STATUS:", response.status_code)

    if response.status_code != 200:
        print("X API RESPONSE:", response.text)
        raise Exception(
            f"X API request failed: {response.status_code}"
        )

    result = response.json()

    users = {
        user["id"]: {
            "username": user.get("username", "Unknown"),
            "profile_image_url": user.get("profile_image_url", "")
        }
        for user in result.get(
            "includes",
            {}
        ).get(
            "users",
            []
        )
    }

    posts = []

    for tweet in result.get("data", []):

        user_info = users.get(
            tweet.get("author_id"),
            {"username": "Unknown", "profile_image_url": ""}
        )

        username = user_info.get("username", "Unknown")

        posts.append({
            "url": (
                f"https://x.com/"
                f"{username}/status/"
                f"{tweet['id']}"
            ),
            "handler_id": f"@{username}",
            "profile_image_url": user_info.get("profile_image_url", ""),
            "created_at": tweet.get(
                "created_at",
                ""
            )
        })

    posts.sort(
        key=lambda x: x.get(
            "created_at",
            ""
        ),
        reverse=True
    )

    return posts[:int(number_of_posts)]


# ============================================================
# GET X USER PROFILE PHOTO
# ============================================================

def get_profile_image_url(username):

    if not TOKEN:
        return ""

    url = f"https://api.x.com/2/users/by/username/{username}"

    response = requests.get(
        url,
        headers={"Authorization": f"Bearer {TOKEN}"},
        params={"user.fields": "profile_image_url"},
        timeout=30
    )

    print("X USER STATUS:", response.status_code)

    if response.status_code != 200:
        print("Could not fetch profile photo:", response.text[:500])
        return ""

    return response.json().get("data", {}).get("profile_image_url", "")


def download_profile_image(image_url, index):

    if not image_url:
        return ""

    try:
        image_url = image_url.replace("_normal.", ".")

        response = requests.get(
            image_url,
            timeout=30,
            headers={
                "User-Agent": "Mozilla/5.0 (compatible; XReportGenerator/1.0)"
            }
        )
        response.raise_for_status()

        filename = os.path.join(
            SCREENSHOT_DIR,
            f"profile_{index}.jpg"
        )

        with open(filename, "wb") as file:
            file.write(response.content)

        return filename

    except Exception as e:
        print("Profile photo download failed:", str(e))
        return ""


# ============================================================
# DIRECT X POST URL
# ============================================================

def get_post_from_url(url):

    pattern = (
        r"(?:https?://)?"
        r"(?:www\.)?"
        r"(?:x\.com|twitter\.com)/"
        r"([^/]+)/status/(\d+)"
    )

    match = re.search(
        pattern,
        url
    )

    if not match:
        raise Exception(
            "Invalid X post URL."
        )

    username = match.group(1)
    post_id = match.group(2)

    return [{
        "url": (
            f"https://x.com/"
            f"{username}/status/"
            f"{post_id}"
        ),
        "handler_id": f"@{username}",
        "profile_image_url": get_profile_image_url(username),
        "created_at": ""
    }]


# ============================================================
# GET OFFICIAL X EMBED HTML
# ============================================================

def get_x_oembed(post_url):

    oembed_url = "https://publish.x.com/oembed"

    params = {
        "url": post_url,
        "maxwidth": 550,
        "hide_thread": "true",
        "omit_script": "false",
        "lang": "en",
        "theme": "light"
    }

    print("Requesting official X oEmbed...")

    response = requests.get(
        oembed_url,
        params=params,
        timeout=30,
        headers={
            "User-Agent": (
                "Mozilla/5.0 "
                "(compatible; XReportGenerator/1.0)"
            )
        }
    )

    print(
        "X oEmbed STATUS:",
        response.status_code
    )

    if response.status_code != 200:
        print(
            "X oEmbed RESPONSE:",
            response.text[:1000]
        )

        raise Exception(
            "X oEmbed request failed: "
            f"{response.status_code}"
        )

    data = response.json()

    html = data.get("html")

    if not html:
        raise Exception(
            "X oEmbed returned no embed HTML."
        )

    print(
        "X oEmbed HTML received."
    )

    return html


# ============================================================
# CREATE SCREENSHOT PAGE
# ============================================================

def create_embed_page_html(embed_html, post_url):

    safe_url = escape(post_url, quote=True)

    return f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>X Post Screenshot</title>

    <style>
        * {{
            box-sizing: border-box;
        }}

        html,
        body {{
            margin: 0;
            padding: 0;
            background: #ffffff;
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Helvetica, Arial, sans-serif;
        }}

        .browser-window {{
            width: 580px;
            margin: 0 auto;
            background: #ffffff;
            border: 1px solid #dadce0;
            border-radius: 8px;
            overflow: hidden;
            box-shadow: 0 4px 12px rgba(0,0,0,0.08);
        }}

        .chrome-tab-bar {{
            height: 38px;
            background: #dee1e6;
            display: flex;
            align-items: flex-end;
            padding-left: 8px;
            padding-right: 8px;
        }}

        .chrome-tab {{
            height: 32px;
            width: 200px;
            background: #ffffff;
            border-radius: 8px 8px 0 0;
            display: flex;
            align-items: center;
            padding: 0 12px;
            gap: 8px;
            font-size: 12px;
            color: #3c4043;
            font-weight: 500;
        }}

        .chrome-tab-icon {{
            width: 14px;
            height: 14px;
            background: #000000;
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            color: #fff;
            font-size: 9px;
            font-weight: bold;
        }}

        .chrome-tab-title {{
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
            flex: 1;
        }}

        .chrome-toolbar {{
            height: 42px;
            background: #ffffff;
            border-bottom: 1px solid #e0e0e0;
            display: flex;
            align-items: center;
            padding: 0 8px;
            gap: 8px;
        }}

        .nav-btn {{
            width: 26px;
            height: 26px;
            display: flex;
            align-items: center;
            justify-content: center;
            color: #5f6368;
            font-size: 14px;
            cursor: pointer;
        }}

        .chrome-address-bar {{
            flex: 1;
            height: 28px;
            background: #f1f3f4;
            border-radius: 14px;
            display: flex;
            align-items: center;
            padding: 0 12px;
            gap: 8px;
            font-size: 13px;
            color: #202124;
        }}

        .lock-icon {{
            color: #5f6368;
            font-size: 12px;
        }}

        .url-text {{
            white-space: nowrap;
            overflow: hidden;
            text-overflow: ellipsis;
        }}

        .browser-content {{
            width: 100%;
            background: #ffffff;
            padding: 0;
            margin: 0;
            display: flex;
            justify-content: center;
        }}

        #tweet-container {{
            width: 100%;
            margin: 0;
            padding: 0;
            background: #ffffff;
            display: flex;
            justify-content: center;
        }}

        blockquote.twitter-tweet {{
            margin: 0 auto !important;
        }}

        iframe.twitter-tweet-rendered {{
            margin: 0 auto !important;
        }}
    </style>
</head>

<body>
    <div class="browser-window">
        <div class="chrome-tab-bar">
            <div class="chrome-tab">
                <div class="chrome-tab-icon">X</div>
                <div class="chrome-tab-title">X Post</div>
            </div>
        </div>

        <div class="chrome-toolbar">
            <div class="nav-btn">←</div>
            <div class="nav-btn">→</div>
            <div class="nav-btn">↻</div>
            <div class="chrome-address-bar">
                <span class="lock-icon">🔒</span>
                <span class="url-text">{safe_url}</span>
            </div>
            <div class="nav-btn">⋮</div>
        </div>

        <div class="browser-content">
            <div id="tweet-container">
                {embed_html}
            </div>
        </div>
    </div>

    <script async src="https://platform.x.com/widgets.js" charset="utf-8"></script>
</body>
</html>
"""


# ============================================================
# WAIT FOR X EMBED
# ============================================================

def wait_for_x_embed(page):

    print("Waiting for official X embed to render...")

    try:
        page.locator("blockquote.twitter-tweet").wait_for(
            state="attached",
            timeout=15000
        )
        print("X blockquote found.")
    except Exception:
        print("X blockquote was not found.")

    try:
        page.wait_for_function(
            """
            () => window.twttr && window.twttr.widgets
            """,
            timeout=20000
        )
        print("X widgets.js loaded.")
    except Exception:
        print("X widgets.js did not become available in time.")

    try:
        page.evaluate(
            """
            () => {
                if (window.twttr && window.twttr.widgets) {
                    window.twttr.widgets.load(
                        document.getElementById('tweet-container')
                    );
                }
            }
            """
        )
        print("Requested X widget rendering.")
    except Exception as e:
        print("Could not manually trigger X widget rendering:", str(e))

    try:
        page.wait_for_function(
            """
            () => {
                const iframe = document.querySelector(
                    'iframe.twitter-tweet-rendered, iframe[src*="platform.x.com"], iframe[src*="platform.twitter.com"]'
                );
                if (!iframe) return false;

                const r = iframe.getBoundingClientRect();
                return r.width > 300 && r.height > 150;
            }
            """,
            timeout=30000
        )
        print("Rendered X tweet iframe found.")
    except Exception:
        print("Rendered X tweet iframe was not found.")
        return False

    page.wait_for_timeout(2000)

    return True


# ============================================================
# TAKE REAL X EMBED SCREENSHOTS (CROPPED AT TIMESTAMP)
# ============================================================

def take_screenshots(posts):

    successful_posts = []

    with sync_playwright() as p:

        print("Starting Chromium...")

        browser = p.chromium.launch(
            headless=True,
            args=[
                "--no-sandbox",
                "--disable-setuid-sandbox",
                "--disable-dev-shm-usage",
                "--disable-gpu"
            ]
        )

        context = browser.new_context(
            viewport={
                "width": 1280,
                "height": 1000
            },
            locale="en-US",
            timezone_id="Asia/Karachi",
            color_scheme="light"
        )

        context.set_default_timeout(15000)

        for i, post in enumerate(posts, 1):

            print(f"Taking screenshot {i}/{len(posts)}...")

            page = None

            try:

                post_url = post["url"]
                print("Post URL:", post_url)

                embed_html = get_x_oembed(post_url)

                html = create_embed_page_html(
                    embed_html,
                    post_url
                )

                page = context.new_page()
                page.set_default_timeout(15000)

                print("Rendering official X embed...")

                page.set_content(
                    html,
                    wait_until="domcontentloaded"
                )

                rendered = wait_for_x_embed(page)

                if not rendered:
                    debug_file = (
                        f"{SCREENSHOT_DIR}/"
                        f"debug_{i}.png"
                    )

                    page.screenshot(
                        path=debug_file,
                        full_page=False
                    )

                    raise Exception(
                        "Official X post did not render. "
                        "Refusing to save a fake/raw-text screenshot."
                    )

                tweet_iframe = page.locator(
                    'iframe.twitter-tweet-rendered, '
                    'iframe[src*="platform.x.com"], '
                    'iframe[src*="platform.twitter.com"]'
                ).first

                tweet_iframe.wait_for(
                    state="visible",
                    timeout=15000
                )

                browser_frame = page.locator(".browser-window")
                browser_box = browser_frame.bounding_box()

                # Get bottom position of the time/date element inside the iframe
                crop_height = page.evaluate(
                    """
                    () => {
                        const iframe = document.querySelector('iframe.twitter-tweet-rendered');
                        if (!iframe || !iframe.contentWindow) return null;
                        
                        const doc = iframe.contentWindow.document;
                        const timeElement = doc.querySelector('time') || doc.querySelector('[aria-label*="PM"], [aria-label*="AM"]');
                        
                        if (timeElement) {
                            const timeRect = timeElement.getBoundingClientRect();
                            const iframeRect = iframe.getBoundingClientRect();
                            const browserRect = document.querySelector('.browser-window').getBoundingClientRect();
                            
                            // Calculate relative offset from browser window top to bottom of timestamp
                            return (iframeRect.top - browserRect.top) + timeRect.bottom + 12;
                        }
                        return null;
                    }
                    """
                )

                filename = (
                    f"{SCREENSHOT_DIR}/"
                    f"post_{i}.png"
                )

                if crop_height and browser_box:
                    # Clip the screenshot directly below the timestamp
                    page.screenshot(
                        path=filename,
                        clip={
                            "x": browser_box["x"],
                            "y": browser_box["y"],
                            "width": browser_box["width"],
                            "height": crop_height
                        }
                    )
                else:
                    # Fallback to full browser window if time calculation fails
                    browser_frame.screenshot(
                        path=filename,
                        timeout=15000
                    )

                if (
                    not os.path.exists(filename)
                    or os.path.getsize(filename) <= 1000
                ):
                    raise Exception(
                        "Screenshot file was not created correctly."
                    )

                post["screenshot"] = filename
                successful_posts.append(post)

                print(
                    f"Screenshot {i} saved successfully."
                )

            except Exception as e:

                print(
                    f"Screenshot {i} failed: {e}"
                )

                traceback.print_exc()

                continue

            finally:

                if page is not None:
                    try:
                        page.close()
                    except Exception:
                        pass

        context.close()
        browser.close()

    return successful_posts


# ============================================================
# CREATE WORD REPORT
# ============================================================

def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcMar = tcPr.first_child_found_in("w:tcMar")

    if tcMar is None:
        from docx.oxml import OxmlElement
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    from docx.oxml.ns import qn

    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{margin}"))
        if node is None:
            from docx.oxml import OxmlElement
            node = OxmlElement(f"w:{margin}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prepare_cell(cell):

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=0, start=60, bottom=0, end=60)

    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Inches(0)
        paragraph.paragraph_format.space_after = Inches(0)
        paragraph.paragraph_format.line_spacing = 1


def create_word_report(posts):

    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    title = document.add_heading(
        "X/Twitter Report",
        level=1
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = document.add_table(
        rows=1,
        cols=5
    )

    table.style = "Table Grid"
    table.autofit = False

    headers = [
        "Sr.No",
        "Screenshot",
        "Post URL",
        "Handler ID",
        "Profile Photo"
    ]

    widths = [
        0.45,
        2.65,
        1.85,
        0.80,
        0.75
    ]

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.width = Inches(widths[i])
        prepare_cell(cell)

        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for number, post in enumerate(posts, 1):

        profile_path = download_profile_image(
            post.get("profile_image_url", ""),
            number
        )

        post["profile_image"] = profile_path

        cells = table.add_row().cells

        for i, width in enumerate(widths):
            cells[i].width = Inches(width)
            prepare_cell(cells[i])

        cells[0].text = str(number)

        screenshot_paragraph = cells[1].paragraphs[0]
        screenshot_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        screenshot_run = screenshot_paragraph.add_run()
        screenshot_run.add_picture(
            post["screenshot"],
            width=Inches(2.60)
        )

        cells[2].text = post["url"]
        cells[3].text = post["handler_id"]

        profile_paragraph = cells[4].paragraphs[0]
        profile_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if profile_path and os.path.exists(profile_path):
            profile_run = profile_paragraph.add_run()
            profile_run.add_picture(
                profile_path,
                width=Inches(0.48)
            )
        else:
            profile_paragraph.add_run("N/A")

        for cell in cells:
            prepare_cell(cell)

    output_file = (
        f"{REPORT_DIR}/X_Report.docx"
    )

    document.save(output_file)

    print(
        "Report saved:",
        output_file
    )

    return output_file


# ============================================================
# LOGIN
# ============================================================

LOGIN_PAGE = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>X Report Generator - Login</title>
    <style>
        * { box-sizing: border-box; }
        body {
            margin: 0;
            min-height: 100vh;
            display: flex;
            align-items: center;
            justify-content: center;
            background: #f5f7fa;
            font-family: Arial, Helvetica, sans-serif;
        }
        .login-box {
            width: 360px;
            background: #ffffff;
            border: 1px solid #e1e5ea;
            border-radius: 12px;
            padding: 30px;
            box-shadow: 0 8px 30px rgba(0,0,0,0.08);
        }
        h2 { margin: 0 0 8px; text-align: center; }
        p { margin: 0 0 22px; text-align: center; color: #666; }
        input {
            width: 100%;
            height: 44px;
            padding: 0 12px;
            border: 1px solid #cfd5dc;
            border-radius: 8px;
            font-size: 15px;
        }
        button {
            width: 100%;
            height: 44px;
            margin-top: 14px;
            border: 0;
            border-radius: 8px;
            background: #111827;
            color: white;
            font-size: 15px;
            cursor: pointer;
        }
        .error { color: #c62828; text-align: center; margin-top: 12px; }
    </style>
</head>
<body>
    <div class="login-box">
        <h2>X/Twitter Report Generator</h2>
        <p>Enter password to continue</p>
        <form method="POST">
            <input type="password" name="password" placeholder="Password" required autofocus>
            <button type="submit">Login</button>
        </form>
        {% if error %}<div class="error">{{ error }}</div>{% endif %}
    </div>
</body>
</html>
"""


@app.route("/login", methods=["GET", "POST"])
def login():

    if request.method == "POST":
        password = request.form.get("password", "")

        if password == LOGIN_PASSWORD:
            session["logged_in"] = True
            return redirect(url_for("home"))

        return LOGIN_PAGE.replace(
            "{% if error %}<div class=\"error\">{{ error }}</div>{% endif %}",
            '<div class="error">Incorrect password.</div>'
        )

    return LOGIN_PAGE.replace(
        "{% if error %}<div class=\"error\">{{ error }}</div>{% endif %}",
        ""
    )


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


def require_login():
    return session.get("logged_in") is True


# ============================================================
# HOME PAGE
# ============================================================

@app.route("/")
def home():

    if not require_login():
        return redirect(url_for("login"))

    return render_template(
        "index.html"
    )


# ============================================================
# GENERATE REPORT
# ============================================================

@app.route(
    "/generate",
    methods=["POST"]
)
def generate():

    if not require_login():
        return jsonify({
            "error": "Login required."
        }), 401

    try:

        data = request.get_json(
            silent=True
        )

        if not data:

            return jsonify({
                "error":
                    "Invalid request."
            }), 400

        user_input = data.get(
            "input",
            ""
        ).strip()

        input_type = data.get(
            "input_type",
            "search"
        )

        number_of_posts = int(
            data.get(
                "number_of_posts",
                10
            )
        )

        if not user_input:

            return jsonify({
                "error":
                    "Please enter something."
            }), 400

        if input_type == "search":

            number_of_posts = max(
                1,
                min(
                    number_of_posts,
                    100
                )
            )

        print("")
        print(
            "=============================="
        )

        print(
            "NEW REPORT REQUEST"
        )

        print(
            "=============================="
        )

        if input_type == "url":

            print(
                "Processing direct X URL..."
            )

            posts = get_post_from_url(
                user_input
            )

        else:

            print(
                "Searching X..."
            )

            print(
                "Requested posts:",
                number_of_posts
            )

            posts = search_posts(
                user_input,
                number_of_posts
            )

        print(
            "Posts found:",
            len(posts)
        )

        if not posts:

            return jsonify({
                "error":
                    "No posts found."
            }), 404

        print(
            "Starting screenshot process..."
        )

        posts = take_screenshots(
            posts
        )

        print(
            "Successful screenshots:",
            len(posts)
        )

        if not posts:

            return jsonify({
                "error":
                    "Could not capture the X posts."
            }), 500

        print(
            "Creating Word report..."
        )

        create_word_report(
            posts
        )

        print(
            "REPORT COMPLETE"
        )

        print(
            "=============================="
        )

        return jsonify({

            "success": True,

            "download_url":
                "/download",

            "posts":
                len(posts)

        })

    except Exception as e:

        print("")
        print(
            "!!!!!!!!!!!!!!!!!!!!!!!!!!!!"
        )

        print(
            "ERROR WHILE GENERATING REPORT"
        )

        print(
            "!!!!!!!!!!!!!!!!!!!!!!!!!!!!"
        )

        print(
            "ERROR:",
            str(e)
        )

        traceback.print_exc()

        print(
            "!!!!!!!!!!!!!!!!!!!!!!!!!!!!"
        )

        print("")

        return jsonify({
            "error":
                "Something went wrong."
        }), 500


# ============================================================
# DOWNLOAD REPORT
# ============================================================

@app.route("/download")
def download():

    if not require_login():
        return jsonify({
            "error": "Login required."
        }), 401

    file_path = (
        f"{REPORT_DIR}/X_Report.docx"
    )

    if not os.path.exists(
        file_path
    ):

        return jsonify({
            "error":
                "Report not found."
        }), 404

    return send_file(
        file_path,
        as_attachment=True,
        download_name="X_Report.docx"
    )


# ============================================================
# START SERVER
# ============================================================

if __name__ == "__main__":

    print("")
    print(
        "======================================"
    )

    print(
        "X/Twitter Report Generator"
    )

    print(
        "Starting Flask server..."
    )

    print(
        "======================================"
    )

    print("")

    app.run(
        host="127.0.0.1",
        port=5000,
        debug=False,
        use_reloader=False
    )