from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import json
import os


def create_report(posts):
    document = Document()

    document.add_heading("X/Twitter Report", level=1)

    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    headers = [
        "Serial No.",
        "Screenshot",
        "Post URL",
        "Handler ID"
    ]

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for number, post in enumerate(posts, 1):
        row = table.add_row().cells

        # Serial number
        row[0].text = str(number)

        # Screenshot
        paragraph = row[1].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if os.path.exists(post["screenshot"]):
            run = paragraph.add_run()
            run.add_picture(
                post["screenshot"],
                width=Inches(2.3)
            )

        # URL
        row[2].text = post["url"]

        # Handler ID
        row[3].text = post["handler_id"]

        for cell in row:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    document.save("X_Report.docx")

    print("Report created successfully: X_Report.docx")


if __name__ == "__main__":

    # Read the posts saved by capture_posts.py
    with open("posts.json", "r", encoding="utf-8") as file:
        posts = json.load(file)

    create_report(posts)