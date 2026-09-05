import os
import re
import traceback
import requests

from flask import Flask, render_template, request, jsonify, send_file
from dotenv import load_dotenv

from playwright.sync_api import sync_playwright

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


# ============================================================
# CONFIGURATION
# ============================================================

load_dotenv()

TOKEN = os.getenv("X_BEARER_TOKEN")

app = Flask(__name__)

SCREENSHOT_DIR = "screenshots"
REPORT_DIR = "reports"

os.makedirs(SCREENSHOT_DIR, exist_ok=True)
os.makedirs(REPORT_DIR, exist_ok=True)


# ============================================================
# SEARCH X POSTS
# ============================================================

def search_posts(query, number_of_posts):

    if not TOKEN:
        raise Exception("X_BEARER_TOKEN is missing.")

    url = "https://api.x.com/2/tweets/search/recent"

    headers = {
        "Authorization": f"Bearer {TOKEN}"
    }

    params = {
        "query": f"{query} -is:retweet",
        "max_results": max(
            10,
            min(int(number_of_posts), 100)
        ),
        "tweet.fields": "created_at,author_id",
        "expansions": "author_id",
        "user.fields": "username"
    }

    response = requests.get(
        url,
        headers=headers,
        params=params,
        timeout=30
    )

    print("X API STATUS:", response.status_code)

    if response.status_code != 200:
        print("X API RESPONSE:", response.text)
        raise Exception(
            f"X API request failed: {response.status_code}"
        )

    result = response.json()

    users = {
        user["id"]: user["username"]
        for user in result.get(
            "includes",
            {}
        ).get(
            "users",
            []
        )
    }

    posts = []

    for tweet in result.get("data", []):

        username = users.get(
            tweet.get("author_id"),
            "Unknown"
        )

        posts.append({
            "url": (
                f"https://x.com/"
                f"{username}/status/"
                f"{tweet['id']}"
            ),
            "handler_id": f"@{username}",
            "created_at": tweet.get(
                "created_at",
                ""
            )
        })

    # Latest posts first
    posts.sort(
        key=lambda x: x.get(
            "created_at",
            ""
        ),
        reverse=True
    )

    return posts[:int(number_of_posts)]


# ============================================================
# DIRECT X POST URL
# ============================================================

def get_post_from_url(url):

    pattern = (
        r"(?:https?://)?"
        r"(?:www\.)?"
        r"(?:x\.com|twitter\.com)/"
        r"([^/]+)/status/(\d+)"
    )

    match = re.search(
        pattern,
        url
    )

    if not match:
        raise Exception(
            "Invalid X post URL."
        )

    username = match.group(1)
    post_id = match.group(2)

    return [{
        "url": (
            f"https://x.com/"
            f"{username}/status/"
            f"{post_id}"
        ),
        "handler_id": f"@{username}",
        "created_at": ""
    }]


# ============================================================
# FIND ORIGINAL POST
# ============================================================

def find_original_post(page, post_id):

    selectors = [
        'article[data-testid="tweet"]',
        'article[role="article"]',
        'article'
    ]

    # --------------------------------------------------------
    # First attempt
    # --------------------------------------------------------

    for selector in selectors:

        try:

            locator = page.locator(selector)

            count = locator.count()

            print(
                f"{selector}: {count}"
            )

            if count == 0:
                continue

            # The first article on a direct status page
            # should normally be the requested post.
            for index in range(
                min(count, 10)
            ):

                candidate = locator.nth(index)

                try:

                    candidate.wait_for(
                        state="visible",
                        timeout=3000
                    )

                    box = candidate.bounding_box()

                    if box is None:
                        continue

                    # Check whether this article contains
                    # the requested status URL.
                    links = candidate.locator(
                        f'a[href*="/status/{post_id}"]'
                    )

                    if links.count() > 0:
                        print(
                            "Found exact original post."
                        )

                        return candidate

                except Exception:
                    continue

            # If exact status link wasn't found,
            # use first visible article.
            for index in range(
                min(count, 10)
            ):

                candidate = locator.nth(index)

                try:

                    candidate.wait_for(
                        state="visible",
                        timeout=2000
                    )

                    if candidate.bounding_box() is not None:

                        print(
                            "Using first visible article."
                        )

                        return candidate

                except Exception:
                    continue

        except Exception:
            continue

    return None


# ============================================================
# TAKE POST SCREENSHOTS
# ============================================================

def take_screenshots(posts):

    successful_posts = []

    with sync_playwright() as p:

        print("Starting Chromium...")

        browser = p.chromium.launch(
            headless=True,
            args=[
                "--no-sandbox",
                "--disable-setuid-sandbox",
                "--disable-dev-shm-usage",
                "--disable-gpu",
                "--disable-blink-features=AutomationControlled"
            ]
        )

        context = browser.new_context(

            viewport={
                "width": 1400,
                "height": 1200
            },

            user_agent=(
                "Mozilla/5.0 "
                "(Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 "
                "(KHTML, like Gecko) "
                "Chrome/131.0.0.0 "
                "Safari/537.36"
            ),

            locale="en-US",

            timezone_id="UTC"
        )

        page = context.new_page()

        # Reduce obvious automation signals
        page.add_init_script("""
            Object.defineProperty(
                navigator,
                'webdriver',
                {
                    get: () => undefined
                }
            );
        """)

        for i, post in enumerate(
            posts,
            1
        ):

            print(
                f"Taking screenshot "
                f"{i}/{len(posts)}..."
            )

            try:

                # ------------------------------------------------
                # Extract post ID
                # ------------------------------------------------

                match = re.search(
                    r"/status/(\d+)",
                    post["url"]
                )

                if not match:
                    print(
                        f"Invalid post URL for post {i}"
                    )
                    continue

                post_id = match.group(1)

                # ------------------------------------------------
                # Open X post
                # ------------------------------------------------

                print(
                    "Opening:",
                    post["url"]
                )

                response = page.goto(
                    post["url"],
                    wait_until="domcontentloaded",
                    timeout=60000
                )

                if response:
                    print(
                        "Page status:",
                        response.status
                    )

                print(
                    "Page title:",
                    page.title()
                )

                # ------------------------------------------------
                # Wait for JavaScript
                # ------------------------------------------------

                page.wait_for_timeout(10000)

                # ------------------------------------------------
                # Scroll slightly to trigger rendering
                # ------------------------------------------------

                try:
                    page.evaluate(
                        "window.scrollTo(0, 300)"
                    )

                    page.wait_for_timeout(3000)

                    page.evaluate(
                        "window.scrollTo(0, 0)"
                    )

                    page.wait_for_timeout(3000)

                except Exception:
                    pass

                # ------------------------------------------------
                # Detect common X login/interstitial pages
                # ------------------------------------------------

                body_text = ""

                try:
                    body_text = page.locator(
                        "body"
                    ).inner_text(
                        timeout=5000
                    )

                except Exception:
                    pass

                body_lower = body_text.lower()

                if (
                    "sign in to x" in body_lower
                    or "sign in" in body_lower
                    and "sign up" in body_lower
                ):

                    print(
                        "X returned a login/interstitial page."
                    )

                # ------------------------------------------------
                # Find original post
                # ------------------------------------------------

                print(
                    "Looking for original post..."
                )

                tweet = find_original_post(
                    page,
                    post_id
                )

                # ------------------------------------------------
                # Retry after additional wait
                # ------------------------------------------------

                if tweet is None:

                    print(
                        "Post not found yet. "
                        "Waiting longer..."
                    )

                    page.wait_for_timeout(
                        10000
                    )

                    tweet = find_original_post(
                        page,
                        post_id
                    )

                # ------------------------------------------------
                # Still not found
                # ------------------------------------------------

                if tweet is None:

                    print(
                        f"Could not locate post {i}"
                    )

                    debug_file = (
                        f"{SCREENSHOT_DIR}/"
                        f"debug_{i}.png"
                    )

                    page.screenshot(
                        path=debug_file,
                        full_page=False
                    )

                    # Also save page HTML for debugging
                    try:

                        html_file = (
                            f"{SCREENSHOT_DIR}/"
                            f"debug_{i}.html"
                        )

                        with open(
                            html_file,
                            "w",
                            encoding="utf-8"
                        ) as file:

                            file.write(
                                page.content()
                            )

                    except Exception:
                        pass

                    continue

                # ------------------------------------------------
                # Screenshot ONLY the post article
                # ------------------------------------------------

                filename = (
                    f"{SCREENSHOT_DIR}/"
                    f"post_{i}.png"
                )

                tweet.screenshot(
                    path=filename
                )

                post["screenshot"] = filename

                successful_posts.append(
                    post
                )

                print(
                    f"Screenshot {i} saved."
                )

            except Exception as e:

                print(
                    f"Screenshot {i} failed:"
                )

                print(
                    str(e)
                )

                traceback.print_exc()

        context.close()
        browser.close()

    return successful_posts


# ============================================================
# CREATE WORD REPORT
# ============================================================

def create_word_report(posts):

    document = Document()

    title = document.add_heading(
        "X/Twitter Report",
        level=1
    )

    title.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    table = document.add_table(
        rows=1,
        cols=4
    )

    table.style = "Table Grid"

    headers = [
        "Sr.No",
        "Screenshot",
        "Post URL",
        "Handler ID"
    ]

    # --------------------------------------------------------
    # Header
    # --------------------------------------------------------

    for i, header in enumerate(headers):

        cell = (
            table.rows[0]
            .cells[i]
        )

        cell.text = header

        cell.vertical_alignment = (
            WD_CELL_VERTICAL_ALIGNMENT.CENTER
        )

        for paragraph in cell.paragraphs:

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
            )

            for run in paragraph.runs:

                run.bold = True

    # --------------------------------------------------------
    # Rows
    # --------------------------------------------------------

    for number, post in enumerate(
        posts,
        1
    ):

        cells = (
            table
            .add_row()
            .cells
        )

        # Sr.No
        cells[0].text = str(number)

        # Screenshot
        paragraph = (
            cells[1]
            .paragraphs[0]
        )

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        run = paragraph.add_run()

        run.add_picture(
            post["screenshot"],
            width=Inches(2.3)
        )

        # Post URL
        cells[2].text = post["url"]

        # Handler ID
        cells[3].text = post["handler_id"]

        for cell in cells:

            cell.vertical_alignment = (
                WD_CELL_VERTICAL_ALIGNMENT.CENTER
            )

    output_file = (
        f"{REPORT_DIR}/X_Report.docx"
    )

    document.save(
        output_file
    )

    print(
        "Report saved:",
        output_file
    )

    return output_file


# ============================================================
# HOME PAGE
# ============================================================

@app.route("/")
def home():

    return render_template(
        "index.html"
    )


# ============================================================
# GENERATE REPORT
# ============================================================

@app.route(
    "/generate",
    methods=["POST"]
)
def generate():

    try:

        data = request.get_json(
            silent=True
        )

        if not data:

            return jsonify({
                "error":
                    "Invalid request."
            }), 400

        user_input = data.get(
            "input",
            ""
        ).strip()

        input_type = data.get(
            "input_type",
            "search"
        )

        number_of_posts = int(
            data.get(
                "number_of_posts",
                10
            )
        )

        if not user_input:

            return jsonify({
                "error":
                    "Please enter something."
            }), 400

        if input_type == "search":

            number_of_posts = max(
                1,
                min(
                    number_of_posts,
                    100
                )
            )

        print("")
        print(
            "=============================="
        )

        print(
            "NEW REPORT REQUEST"
        )

        print(
            "=============================="
        )

        # ------------------------------------------------------
        # SEARCH OR DIRECT URL
        # ------------------------------------------------------

        if input_type == "url":

            print(
                "Processing direct X URL..."
            )

            posts = get_post_from_url(
                user_input
            )

        else:

            print(
                "Searching X..."
            )

            print(
                "Requested posts:",
                number_of_posts
            )

            posts = search_posts(
                user_input,
                number_of_posts
            )

        print(
            "Posts found:",
            len(posts)
        )

        if not posts:

            return jsonify({
                "error":
                    "No posts found."
            }), 404

        # ------------------------------------------------------
        # SCREENSHOTS
        # ------------------------------------------------------

        print(
            "Starting screenshot process..."
        )

        posts = take_screenshots(
            posts
        )

        print(
            "Successful screenshots:",
            len(posts)
        )

        if not posts:

            return jsonify({
                "error":
                    "Could not capture the posts."
            }), 500

        # ------------------------------------------------------
        # WORD REPORT
        # ------------------------------------------------------

        print(
            "Creating Word report..."
        )

        report = create_word_report(
            posts
        )

        print(
            "REPORT COMPLETE"
        )

        print(
            "=============================="
        )

        return jsonify({

            "success": True,

            "download_url":
                "/download",

            "posts":
                len(posts)

        })

    except Exception as e:

        print("")
        print(
            "!!!!!!!!!!!!!!!!!!!!!!!!!!!!"
        )

        print(
            "ERROR WHILE GENERATING REPORT"
        )

        print(
            "!!!!!!!!!!!!!!!!!!!!!!!!!!!!"
        )

        print(
            "ERROR:",
            str(e)
        )

        traceback.print_exc()

        print(
            "!!!!!!!!!!!!!!!!!!!!!!!!!!!!"
        )

        print("")

        return jsonify({
            "error":
                "Something went wrong."
        }), 500


# ============================================================
# DOWNLOAD REPORT
# ============================================================

@app.route("/download")
def download():

    file_path = (
        f"{REPORT_DIR}/X_Report.docx"
    )

    if not os.path.exists(
        file_path
    ):

        return jsonify({
            "error":
                "Report not found."
        }), 404

    return send_file(
        file_path,
        as_attachment=True,
        download_name="X_Report.docx"
    )


# ============================================================
# START SERVER
# ============================================================

if __name__ == "__main__":

    print("")
    print(
        "======================================"
    )

    print(
        "X/Twitter Report Generator"
    )

    print(
        "Starting Flask server..."
    )

    print(
        "======================================"
    )

    print("")

    app.run(
        host="127.0.0.1",
        port=5000,
        debug=False,
        use_reloader=False
    )


